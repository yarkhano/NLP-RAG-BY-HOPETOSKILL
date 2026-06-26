"""
harrison_to_docx.py
────────────────────────────────────────────────────────────────────────────
Reads the Harrison's Chapter 126 text file and produces a fully-formatted
Word (.docx) document with:
  • Proper Heading 1 / Heading 2 / Heading 3 styles
  • All tables rebuilt as real Word tables (with header row shading)
  • Bullet / sub-bullet lists
  • Abbreviation footnote boxes
  • Bold inline labels (e.g. "Abbreviation:", "Note:")
  • Times New Roman 12 pt body, 1-inch margins
  • Page numbers in footer

USAGE
─────
  python harrison_to_docx.py input.txt output.docx

  If you omit arguments it defaults to:
    input  → harrison_chapter126.txt   (same folder as this script)
    output → Harrison_Chapter126_Pneumonia.docx
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════

FONT_BODY      = "Times New Roman"
FONT_HEADING   = "Times New Roman"
SIZE_BODY      = 12
SIZE_H1        = 16
SIZE_H2        = 14
SIZE_H3        = 13
SIZE_H4        = 12
COLOR_H1       = "1F3864"   # dark navy
COLOR_H2       = "2E5496"   # medium blue
COLOR_H3       = "404040"   # dark grey
COLOR_TABLE_HD = "1F3864"   # table header bg
COLOR_TABLE_R1 = "EBF0F8"   # alternating row 1
COLOR_TABLE_R2 = "FFFFFF"   # alternating row 2

# ── Known TOP-LEVEL headings (Heading 1) ─────────────────────────────────
H1_PATTERNS = [
    r"^DEFINITION$",
    r"^PATHOPHYSIOLOGY$",
    r"^PATHOLOGY$",
    r"^COMMUNITY-ACQUIRED PNEUMONIA$",
    r"^VENTILATOR-ASSOCIATED PNEUMONIA$",
    r"^HOSPITAL-ACQUIRED PNEUMONIA$",
    r"^GLOBAL IMPACT$",
    r"^TREATMENT$",
    r"^COMPLICATIONS$",
    r"^FOLLOW-UP$",
    r"^PROGNOSIS$",
    r"^PREVENTION$",
]

# ── Known SECTION headings (Heading 2) ───────────────────────────────────
H2_PATTERNS = [
    r"^■?ETIOLOGY$",
    r"^■?EPIDEMIOLOGY$",
    r"^■?CLINICAL MANIFESTATIONS$",
    r"^■?DIAGNOSIS$",
    r"^■?PROGNOSIS$",
    r"^■?PREVENTION.*$",
    r"^■?PATHOGENESIS$",
    r"^ANTIBIOTIC RESISTANCE$",
    r"^INITIAL ANTIBIOTIC MANAGEMENT$",
    r"^ADJUNCTIVE MEASURES$",
    r"^FAILURE TO IMPROVE$",
    r"^SITE OF CARE$",
    r"^EMPIRICAL THERAPY$",
    r"^SPECIFIC TREATMENT$",
    r"^■.*",          # anything starting with ■ is heading 2
]

# ── Known SUB-SECTION headings (Heading 3) ───────────────────────────────
H3_PATTERNS = [
    r"^GRAM'S STAIN AND CULTURE OF SPUTUM$",
    r"^BLOOD CULTURES$",
    r"^URINARY ANTIGEN TESTS$",
    r"^POLYMERASE CHAIN REACTION$",
    r"^SEROLOGY$",
    r"^BIOMARKERS$",
    r"^Clinical Diagnosis$",
    r"^Etiologic Diagnosis$",
    r"^Quantitative-Culture Approach$",
    r"^Clinical Approach$",
    r"^Outpatients$",
    r"^Inpatients$",
    r"^Nonsevere.*$",
    r"^Severe.*$",
    r"^S\. pneumoniae$",
    r"^CA-MRSA$",
    r"^M\. pneumoniae$",
    r"^Gram-Negative Bacilli$",
    r"^Primary Lung Abscesses$",
]

# ═══════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════

def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_page_numbers(doc):
    """Add centered page numbers to every section footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        for tag, ftype in [("begin", None), (None, " PAGE "), ("separate", None), ("end", None)]:
            if tag:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), tag)
                run._r.append(fc)
            else:
                it = OxmlElement("w:instrText")
                it.text = ftype
                run._r.append(it)

def para_fmt(p, space_before=0, space_after=8, line_spacing=None, first_indent=0):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
    if first_indent:
        pf.first_line_indent = Pt(first_indent)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text.lstrip("■").strip())
    run.bold = True
    run.font.size = Pt(SIZE_H1)
    run.font.name = FONT_HEADING
    r, g, b = hex_to_rgb(COLOR_H1)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:color"), COLOR_H1)
    pBdr.append(btm)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text.lstrip("■").strip())
    run.bold = True
    run.font.size = Pt(SIZE_H2)
    run.font.name = FONT_HEADING
    r, g, b = hex_to_rgb(COLOR_H2)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text.strip())
    run.bold = True
    run.italic = True
    run.font.size = Pt(SIZE_H3)
    run.font.name = FONT_HEADING
    r, g, b = hex_to_rgb(COLOR_H3)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, justify=True, first_indent=True):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    # handle inline bold: **word**
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = FONT_BODY
        r.font.size = Pt(SIZE_BODY)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after  = Pt(6)
    pf.space_before = Pt(0)
    if first_indent:
        pf.first_line_indent = Pt(18)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text.lstrip("•·-– ").strip())
    run.font.name = FONT_BODY
    run.font.size = Pt(SIZE_BODY)
    pf = p.paragraph_format
    pf.space_after  = Pt(3)
    pf.space_before = Pt(0)
    left = Pt(360 + level * 360)
    pf.left_indent    = left
    pf.first_line_indent = Pt(-180)
    return p

def add_abbrev_box(doc, text):
    """Render an Abbreviation / Note line in a light-grey shaded paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    # detect label
    m = re.match(r"^(Abbreviations?|Note|aNote)(:?\s*)(.*)", text, re.I | re.S)
    if m:
        rl = p.add_run(m.group(1) + ": ")
        rl.bold = True; rl.font.name = FONT_BODY; rl.font.size = Pt(10)
        rb = p.add_run(m.group(3))
        rb.font.name = FONT_BODY; rb.font.size = Pt(10)
    else:
        rb = p.add_run(text)
        rb.font.name = FONT_BODY; rb.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(0)
    # grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    return p

# ═══════════════════════════════════════════════════════════════════════════
# TABLE BUILDER
# ═══════════════════════════════════════════════════════════════════════════

def build_table(doc, raw_lines):
    """
    Parse a block of lines that look like a table.
    Supports two styles:
      A) Column-header rows followed by data rows separated by whitespace alignment
      B) Two-column KEY/VALUE rows (FACTOR / POSSIBLE PATHOGEN style)
    """
    # Strip blank lines at edges
    while raw_lines and not raw_lines[0].strip():
        raw_lines.pop(0)
    while raw_lines and not raw_lines[-1].strip():
        raw_lines.pop()

    if not raw_lines:
        return

    # ── Detect if it is a pipe-delimited table ──────────────────────────
    pipe_rows = [l for l in raw_lines if "|" in l]
    if len(pipe_rows) >= 2:
        _build_pipe_table(doc, raw_lines)
        return

    # ── Detect if it has a clear two-column layout via large whitespace ──
    # Find the column split point as most-common large gap
    _build_whitespace_table(doc, raw_lines)


def _build_pipe_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*[-|]+\s*$", line):
            continue   # separator row
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c != ""]
        if cells:
            rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad rows
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    col_w = 9360 // n_cols
    col_widths = [col_w] * n_cols

    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.add_row()
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Pt(col_widths[j])
            set_cell_borders(cell)
            if is_header:
                set_cell_bg(cell, COLOR_TABLE_HD)
            elif i % 2 == 0:
                set_cell_bg(cell, COLOR_TABLE_R1)
            else:
                set_cell_bg(cell, COLOR_TABLE_R2)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
            run.bold = is_header
            if is_header:
                r2, g2, b2 = hex_to_rgb("FFFFFF")
                run.font.color.rgb = RGBColor(r2, g2, b2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
    doc.add_paragraph()   # space after table


def _build_whitespace_table(doc, lines):
    """
    Detect column split by finding a consistent large-gap position,
    then render as a 2-column Word table.
    """
    # Find split point: look for runs of 3+ spaces that appear in most lines
    candidates = {}
    for line in lines:
        for m in re.finditer(r"   +", line):
            mid = (m.start() + m.end()) // 2
            key = mid // 5 * 5           # bucket to nearest 5
            candidates[key] = candidates.get(key, 0) + 1

    if not candidates:
        # Just add as body paragraphs
        for line in lines:
            if line.strip():
                add_body(doc, line.strip(), first_indent=False)
        return

    split_pos = max(candidates, key=candidates.get)

    rows = []
    pending_left = ""
    pending_right = ""

    for line in lines:
        if not line.strip():
            if pending_left or pending_right:
                rows.append((pending_left.strip(), pending_right.strip()))
                pending_left = ""
                pending_right = ""
            continue
        # Try to split at our detected position ±10
        lo = max(0, split_pos - 10)
        hi = min(len(line), split_pos + 20)
        gap = re.search(r"  +", line[lo:hi])
        if gap:
            abs_start = lo + gap.start()
            abs_end   = lo + gap.end()
            left_part  = line[:abs_start].strip()
            right_part = line[abs_end:].strip()
            if pending_left and left_part:
                rows.append((pending_left.strip(), pending_right.strip()))
                pending_left  = left_part
                pending_right = right_part
            elif pending_left and not left_part:
                pending_right += " " + right_part
            else:
                pending_left  = left_part
                pending_right = right_part
        else:
            # Whole line is continuation of left column
            if pending_left:
                pending_left += " " + line.strip()
            else:
                pending_left = line.strip()

    if pending_left or pending_right:
        rows.append((pending_left.strip(), pending_right.strip()))

    if not rows:
        return

    # Create 2-column table
    col_widths = [4200, 5160]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (left, right) in enumerate(rows):
        if not left and not right:
            continue
        row = table.add_row()
        is_header = (i == 0 and any(
            kw in left.upper() for kw in
            ["FACTOR", "STATUS", "DISEASE", "NO RISK", "PATHOGEN", "REGIMEN",
             "HOSPITALIZED", "OUTPATIENT", "NON-MDR", "MDR"]
        ))
        for j, txt in enumerate((left, right)):
            cell = row.cells[j]
            cell.width = Pt(col_widths[j])
            set_cell_borders(cell)
            if is_header:
                set_cell_bg(cell, COLOR_TABLE_HD)
            elif i % 2 == 0:
                set_cell_bg(cell, COLOR_TABLE_R1)
            else:
                set_cell_bg(cell, COLOR_TABLE_R2)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
            run.bold = is_header
            if is_header:
                r2, g2, b2 = hex_to_rgb("FFFFFF")
                run.font.color.rgb = RGBColor(r2, g2, b2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# LINE CLASSIFIER
# ═══════════════════════════════════════════════════════════════════════════

def classify_line(line):
    stripped = line.strip()
    if not stripped:
        return "blank"
    # Page number artefact  e.g.  "1009", "1010\nCHAPTER 126" etc.
    if re.match(r"^\d{4}$", stripped):
        return "page_number"
    if re.match(r"^CHAPTER \d+$", stripped):
        return "chapter_label"
    if re.match(r"^PART \d+$", stripped):
        return "part_label"
    # Abbreviation footnote lines
    if re.match(r"^(Abbreviations?|aNote|Note):?\s+", stripped, re.I):
        return "abbrev"
    if re.match(r"^[a-z][\w,\s]+[:;]\s+", stripped) and len(stripped) < 200 and \
       re.match(r"^[a-z]", stripped):
        # footnote continuation lines like "aStrategies with…"
        if stripped[0].islower() and len(stripped) > 10:
            return "footnote"
    # Heading 1
    for pat in H1_PATTERNS:
        if re.match(pat, stripped):
            return "h1"
    # Heading 2
    for pat in H2_PATTERNS:
        if re.match(pat, stripped):
            return "h2"
    # Heading 3
    for pat in H3_PATTERNS:
        if re.match(pat, stripped):
            return "h3"
    # TABLE heading (TABLE 126-x ...)
    if re.match(r"^TABLE \d+-\d+", stripped, re.I):
        return "table_caption"
    # FIGURE heading
    if re.match(r"^FIGURE \d+-\d+", stripped, re.I):
        return "figure_caption"
    # Bullet: lines starting with • or a letter/number bullet
    if re.match(r"^[•·▪–\-]\s+", stripped):
        return "bullet"
    if re.match(r"^\u2022\s+", stripped):
        return "bullet"
    # All-caps short line → possible heading
    if stripped.isupper() and 4 <= len(stripped) <= 80 and " " in stripped:
        return "h2_candidate"
    return "body"


# ═══════════════════════════════════════════════════════════════════════════
# TABLE-BLOCK DETECTOR
# ═══════════════════════════════════════════════════════════════════════════

def looks_like_table_row(line):
    """True if line seems to be a table data row (has large gaps or |)."""
    if "|" in line:
        return True
    # Two or more words separated by 3+ spaces
    if re.search(r"\S   +\S", line):
        return True
    return False


def extract_table_blocks(lines):
    """
    Returns list of (start_idx, end_idx) for blocks that look like tables.
    A table block:  ≥3 consecutive lines that look like table rows,
    optionally preceded by a TABLE caption line.
    """
    blocks = []
    i = 0
    while i < len(lines):
        # Look for table caption
        if classify_line(lines[i]) == "table_caption":
            start = i
            i += 1
            # skip blank
            while i < len(lines) and not lines[i].strip():
                i += 1
            tbl_start = i
            # collect table rows
            while i < len(lines) and (looks_like_table_row(lines[i]) or
                  not lines[i].strip() or
                  classify_line(lines[i]) in ("abbrev", "footnote")):
                i += 1
                # stop if we hit a real heading
                if i < len(lines) and classify_line(lines[i]) in ("h1","h2","h3","table_caption"):
                    break
            if i - tbl_start >= 2:
                blocks.append((start, i))
            continue
        # Look for standalone table blocks (3+ table-like rows without caption)
        if looks_like_table_row(lines[i]) and i+1 < len(lines) and \
           looks_like_table_row(lines[i+1]):
            start = i
            while i < len(lines) and (looks_like_table_row(lines[i]) or
                  (not lines[i].strip() and i+1 < len(lines) and looks_like_table_row(lines[i+1]))):
                i += 1
            if i - start >= 3:
                blocks.append((start, i))
            continue
        i += 1
    return blocks


# ═══════════════════════════════════════════════════════════════════════════
# MAIN CONVERTER
# ═══════════════════════════════════════════════════════════════════════════

def convert(input_path, output_path):
    print(f"Reading  : {input_path}")
    with open(input_path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()

    # ── pre-clean ────────────────────────────────────────────────────────
    # Remove soft-hyphens and page-artefact numbers embedded in words
    raw = raw.replace("\xad", "")
    # Normalize line endings
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n")

    # ── Identify table blocks up-front ────────────────────────────────────
    table_blocks = extract_table_blocks(lines)
    table_ranges = set()
    for s, e in table_blocks:
        for idx in range(s, e):
            table_ranges.add(idx)

    # ── Build document ─────────────────────────────────────────────────────
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)

    # ── Title page ─────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr1 = tp.add_run("CHAPTER 126")
    tr1.bold = True; tr1.font.size = Pt(14); tr1.font.name = FONT_HEADING
    r2, g2, b2 = hex_to_rgb(COLOR_H2)
    tr1.font.color.rgb = RGBColor(r2, g2, b2)

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = tp2.add_run("PNEUMONIA")
    tr2.bold = True; tr2.font.size = Pt(26); tr2.font.name = FONT_HEADING
    r1, g1, b1 = hex_to_rgb(COLOR_H1)
    tr2.font.color.rgb = RGBColor(r1, g1, b1)

    tp3 = doc.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr3 = tp3.add_run("Lionel A. Mandell, Michael S. Niederman")
    tr3.italic = True; tr3.font.size = Pt(13); tr3.font.name = FONT_BODY

    tp4 = doc.add_paragraph()
    tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr4 = tp4.add_run("Harrison's Principles of Internal Medicine")
    tr4.font.size = Pt(12); tr4.font.name = FONT_BODY

    doc.add_paragraph().add_run("")   # spacer
    doc.add_paragraph(style="Normal").add_run("").font.size = Pt(12)

    # Separator line
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(6)
    sep.paragraph_format.space_after  = Pt(18)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "8")
    btm.set(qn("w:color"), COLOR_H1)
    pBdr.append(btm)
    pPr.append(pBdr)

    # ── Process lines ──────────────────────────────────────────────────────
    i = 0
    current_table_caption = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── In a table block? ──────────────────────────────────────────────
        if i in table_ranges:
            kind = classify_line(line)
            if kind == "table_caption":
                # collect entire block
                caption = stripped
                block_end = i
                for s, e in table_blocks:
                    if s == i:
                        block_end = e
                        break
                # caption paragraph
                cp = doc.add_paragraph()
                cr = cp.add_run(caption)
                cr.bold = True; cr.italic = True
                cr.font.name = FONT_HEADING; cr.font.size = Pt(11)
                r2, g2, b2 = hex_to_rgb(COLOR_H2)
                cr.font.color.rgb = RGBColor(r2, g2, b2)
                cp.paragraph_format.space_before = Pt(12)
                cp.paragraph_format.space_after  = Pt(4)
                # collect table data lines
                tbl_lines = []
                j = i + 1
                while j < block_end:
                    tbl_lines.append(lines[j])
                    j += 1
                build_table(doc, tbl_lines)
                i = block_end
                continue
            elif kind in ("abbrev", "footnote"):
                add_abbrev_box(doc, stripped)
                i += 1
                continue
            else:
                # part of a table without caption — collect block
                for s, e in table_blocks:
                    if s == i:
                        tbl_lines = lines[s:e]
                        build_table(doc, tbl_lines)
                        i = e
                        break
                else:
                    i += 1
                continue

        # ── Normal line ────────────────────────────────────────────────────
        kind = classify_line(line)

        if kind in ("blank", "page_number", "chapter_label", "part_label"):
            i += 1
            continue

        if kind == "h1":
            add_h1(doc, stripped)

        elif kind in ("h2", "h2_candidate"):
            add_h2(doc, stripped)

        elif kind == "h3":
            add_h3(doc, stripped)

        elif kind == "table_caption":
            # shouldn't reach here (handled above) but just in case
            cp = doc.add_paragraph()
            cr = cp.add_run(stripped)
            cr.bold = True; cr.italic = True
            cr.font.name = FONT_HEADING; cr.font.size = Pt(11)
            cp.paragraph_format.space_before = Pt(12)
            cp.paragraph_format.space_after  = Pt(4)

        elif kind == "figure_caption":
            fp = doc.add_paragraph()
            fr = fp.add_run(stripped)
            fr.bold = True; fr.font.name = FONT_BODY; fr.font.size = Pt(11)
            fp.paragraph_format.space_before = Pt(8)
            fp.paragraph_format.space_after  = Pt(4)
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif kind == "bullet":
            # detect sub-bullet level by leading spaces
            level = 0
            if line.startswith("    ") or line.startswith("\t"):
                level = 1
            add_bullet(doc, stripped, level=level)

        elif kind in ("abbrev", "footnote"):
            add_abbrev_box(doc, stripped)

        else:
            # body text — try to merge continuation lines
            body_text = stripped
            # look ahead: if next non-blank line is also body, merge
            j = i + 1
            while j < len(lines):
                next_stripped = lines[j].strip()
                if not next_stripped:
                    break
                next_kind = classify_line(lines[j])
                if next_kind not in ("body",):
                    break
                if j in table_ranges:
                    break
                # merge only if current ends without sentence-ending punctuation
                # or next line starts lowercase (continuation)
                if (body_text and body_text[-1] not in ".!?:") or \
                   (next_stripped and next_stripped[0].islower()):
                    body_text += " " + next_stripped
                    j += 1
                else:
                    break
            add_body(doc, body_text)
            i = j
            continue

        i += 1

    # ── Page numbers ───────────────────────────────────────────────────────
    add_page_numbers(doc)

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Saved    : {output_path}")
    print(f"Done ✓")


# ═══════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) >= 3:
        inp  = sys.argv[1]
        outp = sys.argv[2]
    elif len(sys.argv) == 2:
        inp  = sys.argv[1]
        outp = os.path.splitext(inp)[0] + "_formatted.docx"
    else:
        # defaults
        inp  = "s1s.txt"
        outp = "Harrison_Chapter126_Pneumonia.docx"

    if not os.path.exists(inp):
        print(f"ERROR: Input file not found → {inp}")
        print("USAGE: python harrison_to_docx.py  input.txt  output.docx")
        sys.exit(1)

    convert(inp, outp)